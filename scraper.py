import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import time


base_url = "https://docs.pancakeswap.finance/"
visited = set()
to_visit = [base_url]
image_dir = "downloaded_images"

# Create Folder for Images
os.makedirs(image_dir, exist_ok=True)

# Create Document
doc = Document()
doc.add_heading('PancakeSwap Documentation', 0)

def download_image(img_url):
    try:
        response = requests.get(img_url)
        if response.status_code == 200:
            img_name = os.path.basename(img_url.split("?")[0])  # ตัด query string ออก
            img_path = os.path.join(image_dir, img_name)
            with open(img_path, "wb") as f:
                f.write(response.content)
            return img_path
    except Exception as e:
        print(f"Error downloading image {img_url}: {e}")
    return None

while to_visit:
    url = to_visit.pop(0)
    if url in visited:
        continue
    visited.add(url)
    
    try:
        response = requests.get(url)
        if response.status_code != 200:
            print(f"Failed to fetch {url}")
            continue
        
        print(f"Fetched {url}")
        soup = BeautifulSoup(response.text, "html.parser")
        
        # get main content
        content = soup.find("main")
        if content:
            page_title = soup.title.string if soup.title else url
            doc.add_heading(page_title, level=1)
            
            # Loop through elements in the main content
            # and add them to the Word document
            for element in content.find_all(["p", "img", "h1", "h2", "h3", "h4", "li"]):
                if element.name == "p" or element.name == "li":
                    text = element.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text)
                
                elif element.name in ["h1", "h2", "h3", "h4"]:
                    heading_level = int(element.name[1])
                    doc.add_heading(element.get_text(strip=True), level=heading_level)
                
                elif element.name == "img":
                    img_src = element.get("src")
                    if img_src:
                        img_url = urljoin(base_url, img_src)
                        img_path = download_image(img_url)
                        if img_path:
                            try:
                                # Resize image to fit within 5 inches width
                                # and maintain aspect ratio
                                img = Image.open(img_path)
                                width, height = img.size
                                max_width = 5  # inches
                                aspect_ratio = height / width
                                img_width_inch = min(max_width, width / 300)  # สมมติ 300dpi
                                img_height_inch = img_width_inch * aspect_ratio
                                doc.add_picture(img_path, width=Inches(img_width_inch))
                                doc.add_paragraph("(Image from: {})".format(img_url))
                            except Exception as e:
                                print(f"Error inserting image {img_url}: {e}")
        
        # Find all links in the page
        # and add them to the list of URLs to visit
        for a in soup.find_all("a", href=True):
            link = a['href']
            if link.startswith("/"):
                full_link = urljoin(base_url, link)
                if full_link not in visited and base_url in full_link:
                    to_visit.append(full_link)
                    
    except Exception as e:
        print(f"Error fetching {url}: {e}")
    
    time.sleep(0.5)

# save the document
output_file = "output_document.docx"
doc.save(output_file)
print(f"✅ Successfully saved to {output_file}")
